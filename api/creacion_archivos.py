import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import psycopg2
import pandas as pd

class ConsultaBD():
    def __init__(self):
        self.db_config = {
            "host": "localhost",
            "database": "reto_innova",
            "user": "reto_innova2",
            "password": "123"
        }
        self.resultado_consulta = ''
    
    def extraer_datos_cliente_natural(self, condicion):
        consulta = f"""
        SELECT oi.fecha, oi.saldo_total, cn.nombre, 
        cn.apellido, cn.identificacion, p.producto, 
        oi.numero_prestamo, oi.principal_bank_capital, 
        oi.intereses,oi.gastos, oi.feci, oi.comision_fiduciaria, oi.saldo_total

        FROM operacionInicial as oi
        INNER JOIN clientesNaturales as cn ON oi.id_cliente_natural = cn.id_cliente_natural
        INNER JOIN producto as p ON oi.id_producto = p.id_producto
        WHERE cn.id_cliente_natural = {condicion}
 """
        conexion = psycopg2.connect(**self.db_config)
        cursor = conexion.cursor()
        cursor.execute(consulta)
        self.df = pd.read_sql_query(consulta, conexion)
        self.resultado_consulta = self.df.values.tolist()
        self.resultado_consulta[0][0] = self.resultado_consulta[0][0].strftime("%Y-%m-%d")
        cursor.close()
        conexion.close()


class CrearCertificacionSaldosAcreedores(ConsultaBD):
    def __init__(self, condicion) -> bool:
        self.db_config = {
            "host": "localhost",
            "database": "reto_innova",
            "user": "reto_innova2",
            "password": "123"
        }
        self.doc = docx.Document()
        self.extraer_datos_cliente_natural(condicion=0)
        self.crear_certificacion()

    def crear_certificacion(self):
        #Configuracion de la fuente default para todo el archivo

        #Configuracion del titulo
        titulo = self.doc.add_paragraph()
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        titulo_texto = titulo.add_run('CERTIFICACIÓN DE SALDOS ACREEDORES')
        titulo_texto.bold = True
        
        parrafo_largo = self.doc.add_paragraph()
        parrafo_largo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        texto_certificacion = f"""
        Yo, **Nombre del Contador Publico Autorizado**, actuando en mi condición de Contador 
        Público Autorizado de BANESCO (PANAMÁ), S.A, con registro único de contribuyente número 36633-66-264068 D.V. 11, certifico y hago constar que los libros de contabilidad de BANESCO (PANAMA), S.A., arrojan al {str(self.resultado_consulta[0][0])},a favor de dicho Banco por un total de {str(self.resultado_consulta[0][1])}, los siguientes saldos acreedores contra 
        {str(self.resultado_consulta[0][2])+' '+str(self.resultado_consulta[0][3])}, con identificación/Ruc número {str(self.resultado_consulta[0][4])}, correspondiente a las siguientes obligaciones: 
        """
        #Configuracion de la fuente del texto
        parrafo_largo.add_run(texto_certificacion)
        #Doble salto de linea 
        for _ in range(1):
            self.doc.add_paragraph()
        
        num_pres_producto = self.doc.add_paragraph()
        num_pres_producto.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        num_pres_producto.add_run(str(self.resultado_consulta[0][5])+' No. '+str(self.resultado_consulta[0][6]))
        saldo_capital = self.doc.add_paragraph()
        #Se agrega una tabla para formatear mejor el texto
        tabla = self.doc.add_table(rows=5,cols=2)
        tabla.autofit = False
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

        #Celda de Saldo a capital
        celda_izquierda = tabla.cell(0,0)
        celda_izquierda.width = Cm(1.3)
        celda_izquierda.add_paragraph('Saldo a Capital')
        celda_izquierda.aligment = WD_PARAGRAPH_ALIGNMENT.LEFT
        #Celda del valor de saldo a capital
        celda_derecha = tabla.cell(0,1)
        celda_derecha.widht = Cm(1.4)
        celda_derecha.add_paragraph(str(self.resultado_consulta[0][7]))
        celda_derecha.aligment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        celda_izquierda = tabla.cell(1,0)
        celda_izquierda.add_paragraph('Intereses')
        celda_derecha = tabla.cell(1,1)
        celda_derecha.add_paragraph(str(self.resultado_consulta[0][8]))

        celda_izquierda = tabla.cell(2,0)
        celda_izquierda.add_paragraph('Gastos')
        celda_derecha = tabla.cell(2,1)
        celda_derecha.add_paragraph(str(self.resultado_consulta[0][9]))

        celda_izquierda = tabla.cell(3,0)
        celda_izquierda.add_paragraph('Feci')
        celda_derecha = tabla.cell(3,1)
        celda_derecha.add_paragraph(str(self.resultado_consulta[0][10]))

        celda_izquierda = tabla.cell(4,0)
        celda_izquierda.add_paragraph('Comisión Fiduciaria')
        celda_derecha = tabla.cell(4,1)
        celda_derecha.add_paragraph(str(self.resultado_consulta[0][11]))
        
        celda_izquierda = tabla.cell(4,0)
        saldo_total = celda_izquierda.add_paragraph('Saldo Total')
        saldo_total.bold = True
        celda_derecha = tabla.cell(4,1)
        celda_derecha.add_paragraph(str(self.resultado_consulta[0][12]))
        
        for _ in range(4):
            self.doc.add_paragraph()
        texto_justificacion_final = """
        {nombre del apoderado del banco}
        Apoderado
        {identificación del apoderado del banco}
        Revisado y Comprobado Correcto



        {nombre del Contador Público Autorizado},
        {idoneidad del Contador Público Autorizado}
        {identificación del contador público Autorizado}
        Revisado y Comprobado Correcto
        """
        texto_final = self.doc.add_paragraph(texto_justificacion_final)
        texto_final.aligment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.doc.save(' Certificación de Saldo - Préstamo Hipotecario.docx')
              













        





